#!/usr/bin/env python3
import json
import re
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[1]
DIST = ROOT / "dist"
DIST.mkdir(exist_ok=True)
OUT_DOCX = DIST / "taipower_exam_300q_A4_16pt.docx"

QUESTION_FILES = [
    "questions_verified_core_v1.js",
    "questions_confusion_v3.js",
    "questions_verified_high_discrimination_v1.js",
    "questions_verified_high_discrimination_v2.js",
    "questions_verified_high_discrimination_v3.js",
    "questions_verified_high_discrimination_v4.js",
    "questions_verified_high_discrimination_v5.js",
]

UNIT_NAMES = {
    "01": "台灣電力系統概論",
    "02": "電力系統運轉與調度",
    "03": "電力交易市場概述",
    "04": "輔助服務概論",
    "05": "電力交易平台管理規範及作業程序總覽",
    "06": "日前輔助服務市場之參與作法",
    "07": "日前輔助服務市場之交易商品項目規格",
    "08": "日前輔助服務市場之運作",
    "09": "備用容量交易機制",
    "10": "我國電力交易市場推動之方向與展望",
}

TARGET_COUNTS = {
    "01": 20, "02": 25, "03": 25, "04": 25, "05": 35,
    "06": 35, "07": 40, "08": 40, "09": 35, "10": 20,
}


def load_questions():
    """Extract the JSON array passed into .concat(...) from each JS question file."""
    questions = []
    for filename in QUESTION_FILES:
        path = ROOT / filename
        if not path.exists():
            raise FileNotFoundError(path)
        text = path.read_text(encoding="utf-8")
        start = text.find("[")
        end = text.rfind("]")
        if start < 0 or end <= start:
            raise RuntimeError(f"無法解析題庫陣列: {filename}")
        arr = json.loads(text[start:end + 1])
        if not isinstance(arr, list):
            raise RuntimeError(f"題庫格式不是陣列: {filename}")
        for q in arr:
            if not isinstance(q, dict):
                raise RuntimeError(f"題目格式不是物件: {filename}")
            q = dict(q)
            q["_source_file"] = filename
            questions.append(q)
    return questions


def unit_of(q):
    m = re.search(r"單元(\d{2})", str(q.get("tags", "")))
    if not m:
        m = re.search(r"第(\d+)單元", str(q.get("topic", "")))
        return f"{int(m.group(1)):02d}" if m else "99"
    return m.group(1)


def numeric_id(q):
    nums = [int(x) for x in re.findall(r"\d+", str(q.get("id", "")))]
    return tuple(nums) if nums else (9999,)


def is_verified_batch(q):
    tags = str(q.get("tags", ""))
    if "已核對" in tags:
        return True
    # 早期 questions_confusion_v3.js 建立時尚未統一加上「已核對」tag，
    # 但該批題已具官方教材、章節定位與解析，且是既有正式「易混淆」核對批次。
    if q.get("_source_file") == "questions_confusion_v3.js" and "易混淆" in tags:
        return True
    return False


def validate(questions):
    valid = []
    seen = set()
    errors = []
    legacy_verified = 0
    for q in questions:
        qid = str(q.get("id", "")).strip()
        required = [
            "question", "option_a", "option_b", "option_c", "option_d",
            "answer", "explanation", "source_title", "source_locator"
        ]
        if not qid or qid in seen:
            errors.append(f"重複或缺少題號: {qid}")
            continue
        if any(not str(q.get(k, "")).strip() for k in required):
            errors.append(f"欄位不完整: {qid}")
            continue
        if str(q.get("answer", "")).strip().upper() not in "ABCD":
            errors.append(f"答案格式錯誤: {qid}")
            continue
        if not is_verified_batch(q):
            errors.append(f"未通過已核對批次判定: {qid}")
            continue
        unit = unit_of(q)
        if unit not in UNIT_NAMES:
            errors.append(f"無法判定單元: {qid}")
            continue
        if q.get("_source_file") == "questions_confusion_v3.js" and "已核對" not in str(q.get("tags", "")):
            legacy_verified += 1
        seen.add(qid)
        valid.append(q)

    if errors:
        print("題庫守門排除/異常項目:")
        for e in errors:
            print(" -", e)

    counts = Counter(unit_of(q) for q in valid)
    if len(valid) != 300:
        raise RuntimeError(f"已核對有效題數應為300，實際為{len(valid)}；單元分布={dict(sorted(counts.items()))}")
    if dict(sorted(counts.items())) != TARGET_COUNTS:
        raise RuntimeError(f"單元題數不符合300題規格：{dict(sorted(counts.items()))}")

    print(f"Legacy confusion verified batch accepted: {legacy_verified}")
    return sorted(valid, key=lambda q: (int(unit_of(q)), numeric_id(q), str(q.get("id", ""))))


def set_run_font(run, name="Noto Sans CJK TC", size=16, bold=None):
    run.font.name = name
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.rFonts
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rFonts.set(qn(f"w:{attr}"), name)


def add_text_paragraph(doc, text, size=16, bold=False, align=None, before=0, after=0, keep_next=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.keep_with_next = keep_next
    p.paragraph_format.widow_control = True
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)
    return p


def add_page_field(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    set_run_font(run, size=11)


def setup_document():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(1.5)
    sec.bottom_margin = Cm(1.5)
    sec.left_margin = Cm(1.6)
    sec.right_margin = Cm(1.6)
    sec.header_distance = Cm(0.7)
    sec.footer_distance = Cm(0.7)

    normal = doc.styles["Normal"]
    normal.font.name = "Noto Sans CJK TC"
    normal.font.size = Pt(16)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(0)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK TC")

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("台電電力交易平台資格測驗｜300題紙本刷題版　｜　第 ")
    set_run_font(r, size=11)
    add_page_field(footer)
    r2 = footer.add_run(" 頁")
    set_run_font(r2, size=11)
    return doc


def add_cover(doc):
    add_text_paragraph(doc, "台電電力交易平台專業人員資格測驗", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=70, after=14)
    add_text_paragraph(doc, "300題紙本刷題版", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=22)
    add_text_paragraph(doc, "A4｜正文16pt｜題目與答案解析分開", size=16, align=WD_ALIGN_PARAGRAPH.CENTER, after=28)
    add_text_paragraph(doc, "使用方式", size=18, bold=True, after=8)
    notes = [
        "1. 題目冊依第1～10單元排列，請直接在選項前的括號做記號。",
        "2. 請先完成題目冊，再翻至後半部答案解析區對答案。",
        "3. 每題答案解析均保留官方教材名稱與章節定位，方便回查。",
        "4. 本題庫僅納入來源與解析欄位完整、且屬已核對批次的題目。",
        "5. 市場規則可能更新，考前仍應以台電公司最新公告版本為準。",
    ]
    for t in notes:
        add_text_paragraph(doc, t, size=16, after=7)
    add_text_paragraph(doc, "總題數：300題", size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=30, after=8)
    doc.add_page_break()


def add_unit_heading(doc, unit, section_label):
    add_text_paragraph(doc, section_label, size=20, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
    add_text_paragraph(doc, f"第 {int(unit)} 單元｜{UNIT_NAMES[unit]}", size=20, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)


def add_question(doc, number, q):
    level = str(q.get("level", ""))
    topic = str(q.get("topic", ""))
    add_text_paragraph(doc, f"{number}. {q['question']}", size=16, bold=True, before=8, after=3, keep_next=True)
    for idx, letter in enumerate("ABCD"):
        text = str(q[f"option_{letter.lower()}"])
        add_text_paragraph(doc, f"(   ) {letter}. {text}", size=16, after=2, keep_next=(idx < 3))
    meta = doc.add_paragraph()
    meta.paragraph_format.space_before = Pt(2)
    meta.paragraph_format.space_after = Pt(8)
    r = meta.add_run(f"難度：{level}　｜　考點：{topic}")
    set_run_font(r, size=11)


def add_answer(doc, number, q):
    ans = str(q["answer"]).upper()
    add_text_paragraph(doc, f"{number}. 答案：{ans}", size=16, bold=True, before=8, after=3, keep_next=True)
    add_text_paragraph(doc, f"解析：{q['explanation']}", size=16, after=3, keep_next=True)
    add_text_paragraph(doc, f"官方依據：{q['source_title']}｜{q['source_locator']}", size=14, after=8)


def build():
    qs = validate(load_questions())
    numbered = [(i + 1, q) for i, q in enumerate(qs)]
    by_unit = defaultdict(list)
    for n, q in numbered:
        by_unit[unit_of(q)].append((n, q))

    doc = setup_document()
    add_cover(doc)

    add_text_paragraph(doc, "第一部分｜題目冊", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=40, after=18)
    add_text_paragraph(doc, "請先完成本部分，再翻至答案解析。", size=16, align=WD_ALIGN_PARAGRAPH.CENTER, after=20)
    doc.add_page_break()

    for unit in sorted(UNIT_NAMES):
        if unit != "01":
            doc.add_page_break()
        subject = by_unit[unit][0][1].get("subject", "") if by_unit[unit] else ""
        add_unit_heading(doc, unit, subject)
        for n, q in by_unit[unit]:
            add_question(doc, n, q)

    doc.add_page_break()
    add_text_paragraph(doc, "第二部分｜答案與解析", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=40, after=18)
    add_text_paragraph(doc, "建議先完成題目冊後再使用本區。", size=16, align=WD_ALIGN_PARAGRAPH.CENTER, after=20)
    doc.add_page_break()

    for unit in sorted(UNIT_NAMES):
        if unit != "01":
            doc.add_page_break()
        subject = by_unit[unit][0][1].get("subject", "") if by_unit[unit] else ""
        add_unit_heading(doc, unit, subject)
        for n, q in by_unit[unit]:
            add_answer(doc, n, q)

    doc.save(OUT_DOCX)
    counts = Counter(unit_of(q) for q in qs)
    print(f"Created {OUT_DOCX}")
    print("Unit counts:", dict(sorted(counts.items())))
    print("Answer distribution:", dict(sorted(Counter(q['answer'] for q in qs).items())))
    print("Level distribution:", dict(sorted(Counter(str(q.get('level','')) for q in qs).items())))


if __name__ == "__main__":
    build()
