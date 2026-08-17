#!/usr/bin/env python3
import json
import re
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[1]
DIST = ROOT / "dist"
DIST.mkdir(exist_ok=True)
OUT_DOCX = DIST / "taipower_exam_300q_A4_16pt.docx"

QUESTION_FILES = [
    "questions_verified_core_v1.js",
    "questions_confusion_v3.js",
    "questions_verified_high_discrimination_v1.js",
    "questions_verified_high_discrimination_v2.js",
    "questions_verified_high_discrimination_v3.js",
    "questions_verified_high_discrimination_v4.js",
    "questions_verified_high_discrimination_v5.js",
]

UNIT_NAMES = {
    "01": "台灣電力系統概論",
    "02": "電力系統運轉與調度",
    "03": "電力交易市場概述",
    "04": "輔助服務概論",
    "05": "電力交易平台管理規範及作業程序總覽",
    "06": "日前輔助服務市場之參與作法",
    "07": "日前輔助服務市場之交易商品項目規格",
    "08": "日前輔助服務市場之運作",
    "09": "備用容量交易機制",
    "10": "我國電力交易市場推動之方向與展望",
}

TARGET_COUNTS = {
    "01": 20, "02": 25, "03": 25, "04": 25, "05": 35,
    "06": 35, "07": 40, "08": 40, "09": 35, "10": 20,
}


def extract_concat_array(text, filename):
    marker = text.find(".concat")
    if marker < 0:
        raise RuntimeError(f"找不到 .concat 題庫陣列：{filename}")
    start = text.find("[", marker)
    end = text.rfind("]")
    if start < 0 or end <= start:
        raise RuntimeError(f"找不到題庫陣列範圍：{filename}")
    payload = text[start:end + 1]
    try:
        data = json.loads(payload)
    except json.JSONDecodeError as exc:
        raise RuntimeError(f"題庫 JSON 解析失敗：{filename}｜{exc}") from exc
    if not isinstance(data, list):
        raise RuntimeError(f"題庫不是陣列：{filename}")
    return data


def load_questions():
    questions = []
    for filename in QUESTION_FILES:
        path = ROOT / filename
        if not path.exists():
            raise FileNotFoundError(path)
        data = extract_concat_array(path.read_text(encoding="utf-8"), filename)
        for q in data:
            if not isinstance(q, dict):
                raise RuntimeError(f"題目不是物件：{filename}")
            item = dict(q)
            item["_source_file"] = filename
            questions.append(item)
    return questions


def unit_of(q):
    m = re.search(r"單元(\d{2})", str(q.get("tags", "")))
    if m:
        return m.group(1)
    m = re.search(r"第\s*(\d+)\s*單元", str(q.get("topic", "")))
    return f"{int(m.group(1)):02d}" if m else "99"


def numeric_id(q):
    nums = [int(x) for x in re.findall(r"\d+", str(q.get("id", "")))]
    return tuple(nums) if nums else (9999,)


def is_verified(q):
    tags = str(q.get("tags", ""))
    if "已核對" in tags:
        return True
    # 早期易混淆題建立時，來源與解析皆已完成，但尚未統一補上「已核對」標籤。
    return (
        q.get("_source_file") == "questions_confusion_v3.js"
        and "易混淆" in tags
        and bool(str(q.get("source_title", "")).strip())
        and bool(str(q.get("source_locator", "")).strip())
    )


def validate(questions):
    valid = []
    errors = []
    seen = set()
    for q in questions:
        qid = str(q.get("id", "")).strip()
        required = [
            "question", "option_a", "option_b", "option_c", "option_d",
            "answer", "explanation", "source_title", "source_locator"
        ]
        if not qid or qid in seen:
            errors.append(f"重複或缺少題號：{qid}")
            continue
        if any(not str(q.get(k, "")).strip() for k in required):
            errors.append(f"欄位不完整：{qid}")
            continue
        if str(q.get("answer", "")).strip().upper() not in {"A", "B", "C", "D"}:
            errors.append(f"答案格式錯誤：{qid}")
            continue
        if not is_verified(q):
            errors.append(f"未通過已核對批次判定：{qid}")
            continue
        unit = unit_of(q)
        if unit not in UNIT_NAMES:
            errors.append(f"無法判定單元：{qid}")
            continue
        seen.add(qid)
        valid.append(q)

    if errors:
        print("題庫守門排除項目：")
        for e in errors:
            print(" -", e)

    counts = Counter(unit_of(q) for q in valid)
    counts_dict = dict(sorted(counts.items()))
    if len(valid) != 300:
        raise RuntimeError(f"有效題數應為300，實際為{len(valid)}；單元分布={counts_dict}")
    if counts_dict != TARGET_COUNTS:
        raise RuntimeError(f"單元題數不符合規格：{counts_dict}")

    return sorted(valid, key=lambda q: (int(unit_of(q)), numeric_id(q), str(q.get("id", ""))))


def set_run_font(run, size=16, bold=None):
    name = "Noto Sans CJK TC"
    run.font.name = name
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), name)


def add_para(doc, text, size=16, bold=False, align=None, before=0, after=0, keep=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.keep_with_next = keep
    p.paragraph_format.widow_control = True
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)
    return p


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)
    set_run_font(run, size=11)


def setup_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(1.5)
    sec.bottom_margin = Cm(1.5)
    sec.left_margin = Cm(1.6)
    sec.right_margin = Cm(1.6)
    sec.footer_distance = Cm(0.7)

    normal = doc.styles["Normal"]
    normal.font.name = "Noto Sans CJK TC"
    normal.font.size = Pt(16)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK TC")
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(0)

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("台電電力交易平台資格測驗｜300題紙本刷題版　｜　第 ")
    set_run_font(r, size=11)
    add_page_field(footer)
    r2 = footer.add_run(" 頁")
    set_run_font(r2, size=11)
    return doc


def add_cover(doc):
    add_para(doc, "台電電力交易平台專業人員資格測驗", size=24, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, before=70, after=14)
    add_para(doc, "300題紙本刷題版", size=24, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=22)
    add_para(doc, "A4｜正文16pt｜題目與答案解析分開", size=16,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=28)
    add_para(doc, "使用方式", size=18, bold=True, after=8)
    notes = [
        "1. 題目冊依第1～10單元排列，請直接在選項前的括號做記號。",
        "2. 請先完成題目冊，再翻至後半部答案解析區對答案。",
        "3. 每題答案解析均保留官方教材名稱與章節定位，方便回查。",
        "4. 紙本只納入來源、解析與選項欄位完整，且屬已核對批次的題目。",
        "5. 市場規則可能更新，考前仍應以台電公司最新公告版本為準。",
    ]
    for note in notes:
        add_para(doc, note, size=16, after=7)
    add_para(doc, "總題數：300題", size=18, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, before=30, after=8)
    doc.add_page_break()


def add_unit_header(doc, unit, subject):
    add_para(doc, subject, size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=5)
    add_para(doc, f"第 {int(unit)} 單元｜{UNIT_NAMES[unit]}", size=20, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=16)


def add_question(doc, num, q):
    add_para(doc, f"{num}. {q['question']}", size=16, bold=True,
             before=8, after=3, keep=True)
    for i, letter in enumerate("ABCD"):
        add_para(doc, f"(   ) {letter}. {q[f'option_{letter.lower()}']}",
                 size=16, after=2, keep=(i < 3))
    # 輔助資訊不影響閱讀正文，縮小呈現。
    meta = doc.add_paragraph()
    meta.paragraph_format.space_before = Pt(2)
    meta.paragraph_format.space_after = Pt(8)
    r = meta.add_run(f"難度：{q.get('level','')}　｜　考點：{q.get('topic','')}")
    set_run_font(r, size=11)


def add_answer(doc, num, q):
    add_para(doc, f"{num}. 答案：{str(q['answer']).upper()}", size=16, bold=True,
             before=8, after=3, keep=True)
    add_para(doc, f"解析：{q['explanation']}", size=16, after=3, keep=True)
    add_para(doc, f"官方依據：{q['source_title']}｜{q['source_locator']}", size=14, after=8)


def build():
    questions = validate(load_questions())
    numbered = [(i + 1, q) for i, q in enumerate(questions)]
    by_unit = defaultdict(list)
    for num, q in numbered:
        by_unit[unit_of(q)].append((num, q))

    doc = setup_doc()
    add_cover(doc)

    add_para(doc, "第一部分｜題目冊", size=24, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, before=40, after=18)
    add_para(doc, "請先完成本部分，再翻至答案解析。", size=16,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=20)
    doc.add_page_break()

    for idx, unit in enumerate(sorted(UNIT_NAMES)):
        if idx:
            doc.add_page_break()
        subject = by_unit[unit][0][1].get("subject", "")
        add_unit_header(doc, unit, subject)
        for num, q in by_unit[unit]:
            add_question(doc, num, q)

    doc.add_page_break()
    add_para(doc, "第二部分｜答案與解析", size=24, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, before=40, after=18)
    add_para(doc, "建議先完成題目冊後再使用本區。", size=16,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=20)
    doc.add_page_break()

    for idx, unit in enumerate(sorted(UNIT_NAMES)):
        if idx:
            doc.add_page_break()
        subject = by_unit[unit][0][1].get("subject", "")
        add_unit_header(doc, unit, subject)
        for num, q in by_unit[unit]:
            add_answer(doc, num, q)

    doc.save(OUT_DOCX)
    print(f"Created: {OUT_DOCX}")
    print("Unit counts:", dict(sorted(Counter(unit_of(q) for q in questions).items())))
    print("Answer distribution:", dict(sorted(Counter(str(q['answer']).upper() for q in questions).items())))
    print("Level distribution:", dict(sorted(Counter(str(q.get('level','')) for q in questions).items())))


if __name__ == "__main__":
    build()
