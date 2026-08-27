#!/usr/bin/env python3
import argparse
import json
import re
from collections import Counter, defaultdict
from pathlib import Path

import build_print_pdf_v2 as builder
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[1]
DIST = ROOT / "dist"
DIST.mkdir(exist_ok=True)
BANK_JSON = DIST / "question_bank_800.json"
DRAFT_DOCX = DIST / "taipower_exam_800q_A4_14pt_draft.docx"
FINAL_DOCX = DIST / "taipower_exam_800q_A4_14pt.docx"
TOC_JSON = DIST / "toc_pages_800_14pt.json"
EXPECTED_COUNT = 800


def set_run_font(run, size=14, bold=None):
    builder.set_run_font(run, size=size, bold=bold)


def set_page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)
    set_run_font(run, size=9)


def setup_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(1.15)
    sec.bottom_margin = Cm(1.2)
    sec.left_margin = Cm(1.4)
    sec.right_margin = Cm(1.4)
    sec.header_distance = Cm(0.5)
    sec.footer_distance = Cm(0.55)

    normal = doc.styles["Normal"]
    normal.font.name = "Noto Sans CJK TC"
    normal.font.size = Pt(14)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK TC")
    normal.paragraph_format.line_spacing = 1.0
    normal.paragraph_format.space_after = Pt(0)

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("台電電力交易平台資格測驗｜800題紙本刷題版　｜　")
    set_run_font(r, size=9)
    set_page_number(footer)
    return doc


def add_para(doc, text="", size=14, bold=False, align=None, before=0, after=0,
             line_spacing=1.0, keep_next=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.keep_with_next = keep_next
    p.paragraph_format.widow_control = False
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold)
    return p


def add_cover(doc):
    add_para(doc, "台電電力交易平台專業人員資格測驗", size=25, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, before=75, after=14)
    add_para(doc, "800題紙本刷題版", size=23, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    add_para(doc, "A4｜正文 14pt｜緊湊考卷排版", size=14,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=26)
    add_para(doc, "第一部分｜題目冊", size=17, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=5)
    add_para(doc, "第二部分｜答案與解析", size=17, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=32)
    add_para(doc, "使用提醒", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    reminders = [
        "本紙本與網站使用同一套800題正式題庫及守門檢查結果。",
        "先完成題目冊，再翻到後半部答案解析對答案。",
        "題目依第1～10單元排列，方便分單元刷題與複習。",
        "規則類內容可能更新，考前仍應以台電公司最新公告版本為準。",
    ]
    for line in reminders:
        add_para(doc, line, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, after=5)
    add_para(doc, "共 800 題", size=16, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, before=24, after=0)
    doc.add_page_break()


def add_toc_line(doc, label, page):
    p = add_para(doc, size=14, after=2)
    p.paragraph_format.tab_stops.add_tab_stop(
        Cm(17.4), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
    )
    r1 = p.add_run(label)
    set_run_font(r1, size=14)
    r2 = p.add_run(f"\t{page}")
    set_run_font(r2, size=14)


def add_toc(doc, toc_pages=None):
    add_para(doc, "目錄", size=21, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
             before=4, after=12)

    add_para(doc, "第一部分｜題目冊", size=15, bold=True, after=4)
    for unit in sorted(builder.UNIT_NAMES):
        page = "000" if toc_pages is None else toc_pages["questions"][unit]
        add_toc_line(doc, f"第 {int(unit)} 單元　{builder.UNIT_NAMES[unit]}", page)

    add_para(doc, "第二部分｜答案與解析", size=15, bold=True, before=8, after=4)
    for unit in sorted(builder.UNIT_NAMES):
        page = "000" if toc_pages is None else toc_pages["answers"][unit]
        add_toc_line(doc, f"第 {int(unit)} 單元　{builder.UNIT_NAMES[unit]}", page)

    doc.add_page_break()


def add_part_heading(doc, text):
    add_para(doc, text, size=19, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
             before=0, after=8, keep_next=True)


def add_unit_heading(doc, unit, subject, first=False):
    if not first:
        add_para(doc, "", size=14, before=5, after=0)
    add_para(doc, subject, size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
             before=0, after=1, keep_next=True)
    add_para(doc, f"第 {int(unit)} 單元｜{builder.UNIT_NAMES[unit]}", size=17, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=5, keep_next=True)


def add_question(doc, num, q):
    p = add_para(doc, size=14, before=3, after=5, line_spacing=1.0)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(0.65))
    p.paragraph_format.keep_together = True

    rq = p.add_run(f"{num}. {q['question']}")
    set_run_font(rq, size=14, bold=True)
    for letter in "ABCD":
        ro = p.add_run(f"\n\t{letter}. {q[f'option_{letter.lower()}']}")
        set_run_font(ro, size=14, bold=False)


def compact_source(q):
    source = str(q.get("source_title", "")).strip()
    source = re.sub(r"^\d+\.\s*電力交易平台參考資料\s*", "", source)
    return source


def add_answer(doc, num, q):
    p = add_para(doc, size=14, before=3, after=5, line_spacing=1.0)
    p.paragraph_format.keep_together = True
    r1 = p.add_run(f"{num}. 答案：{str(q['answer']).upper()}")
    set_run_font(r1, size=14, bold=True)
    r2 = p.add_run(f"\n解析：{q['explanation']}")
    set_run_font(r2, size=14)
    r3 = p.add_run(f"\n依據：{compact_source(q)}｜{q['source_locator']}")
    set_run_font(r3, size=14)


def unit_of(q):
    m = re.search(r"單元(\d{2})", str(q.get("tags", "")))
    if m:
        return m.group(1)
    return builder.unit_of(q)


def load_exported_questions():
    if not BANK_JSON.exists():
        raise FileNotFoundError(f"找不到正式800題匯出檔：{BANK_JSON}")
    questions = json.loads(BANK_JSON.read_text(encoding="utf-8"))
    if not isinstance(questions, list):
        raise RuntimeError("800題匯出檔格式錯誤：應為陣列")
    if len(questions) != EXPECTED_COUNT:
        raise RuntimeError(f"紙本題庫應為{EXPECTED_COUNT}題，實際{len(questions)}題")

    seen = set()
    for q in questions:
        qid = str(q.get("id", "")).strip()
        if not qid or qid in seen:
            raise RuntimeError(f"題號缺漏或重複：{qid}")
        seen.add(qid)
        answer = str(q.get("answer", "")).upper()
        if answer not in {"A", "B", "C", "D"}:
            raise RuntimeError(f"答案格式錯誤：{qid}")
        for key in ["question", "option_a", "option_b", "option_c", "option_d", "explanation", "source_title", "source_locator"]:
            if not str(q.get(key, "")).strip():
                raise RuntimeError(f"欄位缺漏：{qid}｜{key}")
        if unit_of(q) not in builder.UNIT_NAMES:
            raise RuntimeError(f"無法判定單元：{qid}")
    return questions


def grouped_questions():
    questions = load_exported_questions()
    numbered = [(i + 1, q) for i, q in enumerate(questions)]
    by_unit = defaultdict(list)
    for num, q in numbered:
        by_unit[unit_of(q)].append((num, q))
    missing = [u for u in builder.UNIT_NAMES if not by_unit[u]]
    if missing:
        raise RuntimeError(f"下列單元沒有題目：{missing}")
    return questions, by_unit


def build_docx(output_path, toc_pages=None):
    questions, by_unit = grouped_questions()
    doc = setup_doc()
    add_cover(doc)
    add_toc(doc, toc_pages=toc_pages)

    add_part_heading(doc, "第一部分｜題目冊")
    for idx, unit in enumerate(sorted(builder.UNIT_NAMES)):
        subject = by_unit[unit][0][1].get("subject", "")
        add_unit_heading(doc, unit, subject, first=(idx == 0))
        for num, q in by_unit[unit]:
            add_question(doc, num, q)

    doc.add_page_break()
    add_part_heading(doc, "第二部分｜答案與解析")
    for idx, unit in enumerate(sorted(builder.UNIT_NAMES)):
        subject = by_unit[unit][0][1].get("subject", "")
        add_unit_heading(doc, unit, subject, first=(idx == 0))
        for num, q in by_unit[unit]:
            add_answer(doc, num, q)

    doc.save(output_path)
    print(f"Created: {output_path}")
    print("Unit counts:", dict(sorted(Counter(unit_of(q) for q in questions).items())))
    print("Level counts:", dict(sorted(Counter(str(q.get('level',''))[:1] for q in questions).items())))


def canonical(text):
    return re.sub(r"\s+", "", text).replace("|", "｜")


def extract_toc_pages(pdf_path):
    import fitz

    pdf = fitz.open(pdf_path)
    page_texts = [canonical(page.get_text("text")) for page in pdf]
    result = {"questions": {}, "answers": {}, "page_count": len(pdf)}
    for unit in sorted(builder.UNIT_NAMES):
        needle = canonical(f"第 {int(unit)} 單元｜{builder.UNIT_NAMES[unit]}")
        hits = [idx + 1 for idx, text in enumerate(page_texts) if needle in text]
        if len(hits) < 2:
            raise RuntimeError(f"無法在PDF中找到兩次單元標題：{unit}｜hits={hits}")
        result["questions"][unit] = hits[0]
        result["answers"][unit] = hits[1]
    return result


def verify_pdf(pdf_path, expected):
    import fitz

    actual = extract_toc_pages(pdf_path)
    if actual["questions"] != expected["questions"] or actual["answers"] != expected["answers"]:
        raise RuntimeError(f"最終PDF頁碼與目錄不一致：expected={expected}, actual={actual}")

    pdf = fitz.open(pdf_path)
    all_text = "\n".join(page.get_text("text") for page in pdf)
    answer_count = len(re.findall(r"(?m)^\s*\d+\.\s*答案：", all_text))
    if answer_count != EXPECTED_COUNT:
        raise RuntimeError(f"PDF答案區完整性檢查異常：answers={answer_count}")
    if "800題紙本刷題版" not in canonical(all_text):
        raise RuntimeError("PDF封面未顯示800題紙本刷題版")

    print(f"Verified final PDF: {pdf_path}")
    print(f"Pages: {len(pdf)}")
    print("TOC page map verified.")
    print(f"Answer blocks: {answer_count}")


def main():
    parser = argparse.ArgumentParser()
    mode = parser.add_mutually_exclusive_group(required=True)
    mode.add_argument("--draft", action="store_true")
    mode.add_argument("--final", action="store_true")
    mode.add_argument("--verify", action="store_true")
    parser.add_argument("--draft-pdf")
    parser.add_argument("--pdf")
    args = parser.parse_args()

    if args.draft:
        build_docx(DRAFT_DOCX, toc_pages=None)
        return

    if args.final:
        if not args.draft_pdf:
            raise SystemExit("--final requires --draft-pdf")
        toc_pages = extract_toc_pages(Path(args.draft_pdf))
        TOC_JSON.write_text(json.dumps(toc_pages, ensure_ascii=False, indent=2), encoding="utf-8")
        build_docx(FINAL_DOCX, toc_pages=toc_pages)
        print("TOC:", json.dumps(toc_pages, ensure_ascii=False))
        return

    if args.verify:
        if not args.pdf:
            raise SystemExit("--verify requires --pdf")
        if not TOC_JSON.exists():
            raise SystemExit(f"TOC map missing: {TOC_JSON}")
        expected = json.loads(TOC_JSON.read_text(encoding="utf-8"))
        verify_pdf(Path(args.pdf), expected)


if __name__ == "__main__":
    main()
